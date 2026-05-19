from __future__ import annotations

import html
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOC_IN = ROOT / "docs" / "孙柯涛毕业论文初稿-补充实验图版.docx"
DOC_OUT = ROOT / "docs" / "孙柯涛毕业论文初稿-补充实验图版-第三章补图.docx"
DIAGRAM_DIR = ROOT / "docs" / "diagrams"
FIG_DIR = ROOT / "docs" / "figures" / "ch3_extra"


def cell(
    cell_id: str,
    value: str,
    style: str,
    x: int,
    y: int,
    w: int,
    h: int,
    parent: str = "1",
) -> str:
    return f'''        <mxCell id="{cell_id}" value="{html.escape(value)}" style="{style}" vertex="1" parent="{parent}">
          <mxGeometry x="{x}" y="{y}" width="{w}" height="{h}" as="geometry"/>
        </mxCell>'''


def edge(edge_id: str, source: str, target: str, label: str = "") -> str:
    style = "endArrow=block;html=1;rounded=0;strokeWidth=2;strokeColor=#64748b;"
    return f'''        <mxCell id="{edge_id}" value="{html.escape(label)}" style="{style}" edge="1" parent="1" source="{source}" target="{target}">
          <mxGeometry relative="1" as="geometry"/>
        </mxCell>'''


def wrap_drawio(name: str, inner: str, width: int = 1169, height: int = 827) -> str:
    return f'''<mxfile host="app.diagrams.net" modified="2026-05-06T00:00:00.000Z" agent="5.0" version="21.0.0" type="device">
  <diagram name="{html.escape(name)}" id="{html.escape(name)}">
    <mxGraphModel dx="1422" dy="794" grid="1" gridSize="10" guides="1" tooltips="1" connect="1" arrows="1" fold="1" page="1" pageScale="1" pageWidth="{width}" pageHeight="{height}" math="0" shadow="0">
      <root>
        <mxCell id="0"/>
        <mxCell id="1" parent="0"/>
{inner}
      </root>
    </mxGraphModel>
  </diagram>
</mxfile>
'''


def build_business_flow() -> str:
    title_style = "rounded=0;whiteSpace=wrap;html=1;fillColor=#ffffff;strokeColor=none;fontSize=20;fontStyle=1;fontColor=#1f2933;"
    step_style = "rounded=1;whiteSpace=wrap;html=1;arcSize=8;fillColor=#e0f2fe;strokeColor=#0284c7;fontSize=15;fontStyle=1;fontColor=#0f172a;"
    note_style = "rounded=1;whiteSpace=wrap;html=1;arcSize=8;fillColor=#f8fafc;strokeColor=#cbd5e1;fontSize=13;fontColor=#334155;"
    status_style = "rounded=1;whiteSpace=wrap;html=1;arcSize=8;fillColor=#dcfce7;strokeColor=#16a34a;fontSize=14;fontColor=#14532d;"

    cells = [
        cell("title", "评论分析业务流程", title_style, 430, 30, 300, 42),
        cell("s1", "1. 登录系统", step_style, 60, 130, 150, 68),
        cell("s2", "2. 创建或选择商品", step_style, 250, 130, 170, 68),
        cell("s3", "3. 导入CSV评论", step_style, 470, 130, 170, 68),
        cell("s4", "4. 查看批次统计", step_style, 690, 130, 170, 68),
        cell("s5", "5. 创建分析任务", step_style, 910, 130, 170, 68),
        cell("s6", "6. 执行情感分析", step_style, 910, 285, 170, 68),
        cell("s7", "7. 提取功能点", step_style, 690, 285, 170, 68),
        cell("s8", "8. 挖掘负面问题", step_style, 470, 285, 170, 68),
        cell("s9", "9. 查看总览图表", step_style, 250, 285, 170, 68),
        cell("s10", "10. 生成分析报告", step_style, 60, 285, 170, 68),
        cell("n1", "导入阶段记录总行数、成功数、重复数和失败数，便于判断数据质量。", note_style, 470, 220, 390, 44),
        cell("n2", "分析完成后，情感比例、功能点排行和问题关键词进入总览页与报告页。", status_style, 300, 395, 560, 50),
    ]
    edges = [
        edge("e1", "s1", "s2"),
        edge("e2", "s2", "s3"),
        edge("e3", "s3", "s4"),
        edge("e4", "s4", "s5"),
        edge("e5", "s5", "s6"),
        edge("e6", "s6", "s7"),
        edge("e7", "s7", "s8"),
        edge("e8", "s8", "s9"),
        edge("e9", "s9", "s10"),
    ]
    return wrap_drawio("第三章业务流程图", "\n".join(cells + edges), 1169, 520)


def build_function_structure() -> str:
    title_style = "rounded=0;whiteSpace=wrap;html=1;fillColor=#ffffff;strokeColor=none;fontSize=20;fontStyle=1;fontColor=#1f2933;"
    root_style = "rounded=1;whiteSpace=wrap;html=1;arcSize=8;fillColor=#fef3c7;strokeColor=#d97706;fontSize=16;fontStyle=1;fontColor=#0f172a;"
    group_style = "rounded=1;whiteSpace=wrap;html=1;arcSize=6;fillColor=#eef2ff;strokeColor=#4f46e5;fontSize=15;fontStyle=1;fontColor=#1e1b4b;"
    item_style = "rounded=1;whiteSpace=wrap;html=1;arcSize=6;fillColor=#ffffff;strokeColor=#94a3b8;fontSize=13;fontColor=#334155;align=left;spacingLeft=14;"
    lane_style = "rounded=1;whiteSpace=wrap;html=1;arcSize=4;fillColor=#f8fafc;strokeColor=#cbd5e1;"
    note_style = "rounded=1;whiteSpace=wrap;html=1;arcSize=6;fillColor=#ecfdf5;strokeColor=#16a34a;fontSize=13;fontColor=#14532d;"

    cells = [
        cell("title", "系统功能需求结构", title_style, 440, 22, 310, 42),
        cell("root", "电商评论多维可视化分析平台", root_style, 395, 78, 400, 56),
        cell("lane1", "", lane_style, 45, 170, 330, 295),
        cell("lane2", "", lane_style, 430, 170, 330, 295),
        cell("lane3", "", lane_style, 815, 170, 330, 295),
        cell("g1", "基础数据与权限", group_style, 75, 195, 270, 46),
        cell("g2", "评论分析处理", group_style, 460, 195, 270, 46),
        cell("g3", "结果展示与报告", group_style, 845, 195, 270, 46),
        cell("i11", "用户登录与身份校验", item_style, 75, 265, 270, 38),
        cell("i12", "商品信息新增、编辑、查询", item_style, 75, 315, 270, 38),
        cell("i13", "评论导入与批次统计", item_style, 75, 365, 270, 38),
        cell("i14", "评论有效性管理", item_style, 75, 415, 270, 38),
        cell("i21", "情感三分类识别", item_style, 460, 255, 270, 38),
        cell("i22", "模型路径解析与规则回退", item_style, 460, 305, 270, 38),
        cell("i23", "功能点提取与同义归一", item_style, 460, 355, 270, 38),
        cell("i24", "负面问题关键词挖掘", item_style, 460, 405, 270, 38),
        cell("i31", "分析总览与统计卡片", item_style, 845, 265, 270, 38),
        cell("i32", "情感分布、功能点与问题排行", item_style, 845, 315, 270, 38),
        cell("i33", "任务状态跟踪与重试", item_style, 845, 365, 270, 38),
        cell("i34", "报告记录与PDF导出", item_style, 845, 415, 270, 38),
        cell("note", "三类功能共同支撑“评论导入—模型分析—结果查看—报告生成”的闭环。", note_style, 300, 500, 590, 42),
    ]
    return wrap_drawio("第三章功能需求结构图", "\n".join(cells), 1190, 570)


def export_png(drawio_path: Path, png_path: Path) -> None:
    subprocess.run(
        ["drawio", "-x", "-f", "png", "-o", str(png_path), str(drawio_path)],
        check=True,
        cwd=str(ROOT),
    )


def insert_after(paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if style:
        p.style = style
    if text:
        p.add_run(text)
    return p


def add_picture_after(cursor: Paragraph, image_path: Path, width_cm: float = 14.6) -> Paragraph:
    p = insert_after(cursor, "", "Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Cm(width_cm))
    return p


def add_caption_after(cursor: Paragraph, caption: str) -> Paragraph:
    p = insert_after(cursor, caption, "Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.bold = True
    return p


def set_all_runs_black(document: Document) -> None:
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            color = run._element.get_or_add_rPr().find(qn("w:color"))
            if color is None:
                color = OxmlElement("w:color")
                run._element.get_or_add_rPr().append(color)
            color.set(qn("w:val"), "000000")


def insert_figures(flow_png: Path, structure_png: Path) -> None:
    shutil.copy2(DOC_IN, DOC_OUT)
    doc = Document(DOC_OUT)

    paragraphs = list(doc.paragraphs)
    p_usecase = next(p for p in paragraphs if p.text.strip().startswith("围绕管理员和分析员，系统用例包括"))
    cursor = p_usecase
    cursor = insert_after(cursor, "为说明这些用例在一次实际分析中的先后关系，本文将用户侧业务流程整理为图 3-2。该流程从商品和评论数据准备开始，到分析结果查看和报告生成结束，强调用户在系统中的操作顺序。", "First Paragraph")
    cursor = add_picture_after(cursor, flow_png)
    cursor = add_caption_after(cursor, "图 3-2 评论分析业务流程图")

    paragraphs = list(doc.paragraphs)
    p_requirement = next(p for p in paragraphs if p.text.strip().startswith("结合系统业务流程，系统功能需求分为"))
    p_requirement.text = "结合系统业务流程，系统功能需求可以按基础数据与权限、评论分析处理、结果展示与报告三个层面进行划分，如图 3-3 所示；各功能项的文字说明见表 3-2。"
    cursor = p_requirement
    cursor = add_picture_after(cursor, structure_png)
    cursor = add_caption_after(cursor, "图 3-3 系统功能需求结构图")

    set_all_runs_black(doc)
    doc.save(DOC_OUT)


def main() -> None:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)

    flow_drawio = DIAGRAM_DIR / "08_第三章业务流程图.drawio"
    structure_drawio = DIAGRAM_DIR / "09_第三章功能需求结构图.drawio"
    flow_png = FIG_DIR / "figure_3_2_business_flow.png"
    structure_png = FIG_DIR / "figure_3_3_function_structure.png"

    flow_drawio.write_text(build_business_flow(), encoding="utf-8")
    structure_drawio.write_text(build_function_structure(), encoding="utf-8")
    export_png(flow_drawio, flow_png)
    export_png(structure_drawio, structure_png)
    insert_figures(flow_png, structure_png)

    print(f"drawio: {flow_drawio}")
    print(f"drawio: {structure_drawio}")
    print(f"png: {flow_png}")
    print(f"png: {structure_png}")
    print(f"docx: {DOC_OUT}")


if __name__ == "__main__":
    main()
