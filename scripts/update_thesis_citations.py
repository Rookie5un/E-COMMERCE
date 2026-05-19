#!/usr/bin/env python3
"""Update thesis literature citations and Word cross-links.

This script is intentionally narrow: it only edits the current thesis draft,
rewrites selected paragraphs, rebuilds reference bookmarks, and converts body
citations like [1] into superscript internal hyperlinks to the corresponding
reference item.
"""

from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "docs/基于NLP技术的电商评论多维可视化分析平台_毕业论文_初稿.docx"
REF_COUNT = 20


PARAGRAPH_REPLACEMENTS = {
    "近年来，在线零售已经成为日常消费的重要方式。用户在购物后留下的评论，不仅会影响其他消费者的购买判断，也会给商家提供产品和服务反馈。相比销量、评分、退货率等结构化数据，评论文本能看到更具体的使用体验，例如质量问题、物流问题、售后体验和用户真实感受[1]。":
        "近年来，在线零售已经成为日常消费的重要方式。用户在购物后留下的评论，不只是给其他消费者看的购买参考，也沉淀了商家理解产品和服务的材料。相比销量、评分、退货率等结构化数据，评论文本更容易暴露具体体验，例如质量、物流、售后和使用场景中的真实感受；相关电商数据分析研究也表明，用户行为与评论信息正在成为平台理解消费需求的重要来源[1]。",
    "随着深度学习的发展，CNN、RNN、BiLSTM 和 Attention 等模型开始用于情感分析任务。史振杰等提出 BiLSTM 与 Attention 结合的电商评论情感分析模型，通过双向时序信息和注意力权重提取评论特征[2]。":
        "随着深度学习的发展，CNN、RNN、BiLSTM 和 Attention 等模型逐渐进入评论情感分析。史振杰等在京东手机评论场景中把 BiLSTM 与 Attention 结合，用双向时序信息捕捉前后文，再通过注意力权重突出情感词和关键描述[2]。",
    "这类方法相比词典法能利用更多上下文信息，对评论中的局部重点和语义联系处理得更细。":
        "这类方法的优点是结构清楚，训练和部署成本也可控，适合在中等规模评论数据上做基线实验。",
    "但这类模型仍然较依赖任务数据和词向量质量，在面对中文评论的一词多义、口语化表达和领域迁移时存在一定局限。":
        "不过，中文电商评论里常有省略、转折、同一句正负并存等现象，传统深度模型遇到领域迁移时仍然容易出现误判。",
    "预训练语言模型进一步提升了评论情感分析的效果。陈鹏杰研究了 BERT 在电商评论情感分析中的应用，指出预训练模型有助于捕捉细粒度语义信息[3]。付帅在研究中引入 ERNIE 动态词向量与 RCNN 结构，强调中文评论语境变化对情感分类的影响[4]。":
        "预训练语言模型把评论分析推进了一步。陈鹏杰围绕 BERT 在电商评论中的应用展开研究，强调预训练表示能更好地捕捉上下文和细粒度语义[3]。付帅则将 ERNIE 动态词向量与 RCNN 结构结合，用来缓解中文评论中一词多义和表达不规范带来的影响[4]。",
    "Hou Jin 等面向电商平台用户评论，提出结合 BERT、BiGRU 与 TextCNN 的分析方法，说明预训练表示和下游特征提取可以结合使用[5]。从相关研究看，BERT、RoBERTa、ERNIE 等模型已经常用于中文评论情感分析。":
        "Hou Jin 等面向电商平台用户评论，提出 BERT、BiGRU 与 TextCNN 结合的模型，并使用准确率、精确率和 F1 等指标评价结果[5]。这些研究说明，预训练表示可以作为评论情感识别的底座，再根据具体任务加入序列建模或卷积特征提取。",
    "因此，本文没有只停留在词典规则方案，而是在项目中加入 RoBERTa 情感分类模型的训练和调用流程。":
        "基于这一思路，本文没有把情感分析停留在情感词规则上，而是把 RoBERTa 微调、模型加载和规则兜底都纳入系统流程，既保留模型效果，也保证演示环境下的基本可用性。",
    "评论分析的结果最终需要给用户查看和使用。黄珣等提出 LDA-RoBERTa 双模型融合方案，对在线评论主题和情感进行联合分析[6]。刘梦等讨论了 ECharts 在数据管理平台中的应用，说明图表可以帮助用户更直观地理解统计结果[7]。":
        "评论分析结果最终要能被人看懂。黄珣等将 LDA 主题模型与 RoBERTa 情感分类结合，用京东手机评论分析用户关注主题和情感倾向[6]。刘梦等讨论了 ECharts 在数据管理平台中的应用，强调图表对数据规律呈现和业务理解的作用[7]。这也提醒本文，模型输出如果只停在标签层面，实际使用价值会受到限制。",
    "这类研究给本文的启发是，评论分析不仅要有模型输出，还要把输出结果整理成页面和报告，方便后续查看。":
        "因此，本文在系统设计中把分析任务、统计接口、图表页面和 PDF 报告放在同一条业务链路中，而不是把算法结果单独留在训练日志里。",
    "已有研究表明，预训练语言模型适合处理中文评论情感分类任务，方面级分析和主题归纳也可以补充更细的解释信息[8]。":
        "细粒度评论理解也是本文关注的一个背景。尤良辉、张华熊在丝绸行业电商评论中构建属性-情感分析流程，说明属性级结果比整体情绪更接近业务问题[8]。苏明星等将方面抽取作为商品中文评论分析的重要环节，为功能点识别提供了参考[9]。在数据来源扩展方面，刘多林、吕苗关于 Scrapy 分布式采集的研究可作为后续合规采集模块的技术参照[10]。",
    "但从本科毕业设计的角度看，仅做算法实验或仅做管理页面都不够完整。本文选择把中文电商评论三分类模型训练作为主线，同时实现评论导入、在线分析和可视化展示，使模型训练结果可以在系统中被调用。":
        "近两年的研究进一步说明，这一方向仍在快速演进。胥桂仙等将 BERT 与超图对偶注意力网络结合，尝试增强短文本中的结构信息表达[11]；黄山山基于 RoBERTa 做在线用户评论细粒度情感分析，与本文的模型选型较为接近[12]；金书丞、王嘉梅则从客户评论出发，强调多层语义特征和注意力机制对情感判断的帮助[13]。",
    "预训练语言模型先在大规模语料上进行自监督学习，再迁移到下游任务。BERT 使用双向 Transformer 编码结构，可以根据上下文表示词语含义，弥补传统静态词向量的一些不足[3][9]。":
        "预训练语言模型先在大规模语料上学习通用表示，再迁移到下游任务。BERT 使用双向 Transformer 编码结构，通过 Masked Language Model 等预训练任务获得上下文相关的词表示，后续可以通过微调适配文本分类、问答和序列标注等任务[19]。",
    "RoBERTa 在 BERT 的基础上调整了训练策略和语料使用方式，中文 RoBERTa 模型常用于情感分类和信息抽取任务。本文情感分析模块使用中文 RoBERTa 系列模型进行微调。":
        "RoBERTa 在 BERT 基础上调整了预训练策略，例如更充分的训练数据、更大的 batch 和动态 mask 等，使模型在多类自然语言理解任务上更加稳健[20]。本文情感分析模块使用中文 RoBERTa 系列模型进行微调，原因不是追求复杂结构，而是它在中文评论这种上下文依赖较强的文本上更适合做三分类底座。",
    "项目在情感分析之外增加了功能点提取和负面问题挖掘。AspectExtractor 采用领域词典匹配和归一化方式，把“屏幕、显示、显示屏”等不同说法映射到统一维度。":
        "项目在情感分析之外增加了功能点提取和负面问题挖掘。AspectExtractor 采用领域词典匹配和归一化方式，把“屏幕、显示、显示屏”等不同说法映射到统一维度。属性值提取和方面级情感研究已经证明，更细的属性信息有助于把评论从整体好坏拆解到具体问题上[14][16][17]，但本文当前实现仍以可解释、易维护的词典归一化为主。",
    "训练时使用 AdamW 优化器和线性预热学习率调度器。每轮训练结束后，脚本输出损失、准确率、F1 值和分类报告，并按验证集 macro-F1 保存最佳模型。三分类任务中，各类别数量不完全相同，因此 macro-F1 比单独看准确率更适合观察中性类和负向类表现。":
        "训练时使用 AdamW 优化器和线性预热学习率调度器。每轮训练结束后，脚本输出损失、准确率、F1 值和分类报告，并按验证集 macro-F1 保存最佳模型。Hou Jin 等在电商评论情感分析中也将准确率和 F1 作为主要评价指标[5]；Daza 等的综述进一步说明，不同数据集和类别分布会影响模型比较，因此三分类任务不能只看总体准确率[18]。",
    "最终模型配置可在 backend/data/models/roberta-sentiment-thesis-final/training_args.txt 和 training_summary.json 中查看。模型基座为 hfl/chinese-roberta-wwm-ext，训练数据为 data/train_balanced_full.csv，max_length 为 128，batch_size 为 32，训练 5 轮，学习率为 2e-5，并使用 FGM、早停和交叉熵损失。":
        "最终模型配置可在 backend/data/models/roberta-sentiment-thesis-final/training_args.txt 和 training_summary.json 中查看。模型基座为 hfl/chinese-roberta-wwm-ext，训练数据为 data/train_balanced_full.csv，max_length 为 128，batch_size 为 32，训练 5 轮，学习率为 2e-5，并使用 FGM、早停和交叉熵损失。该配置延续了 BERT/RoBERTa 微调的基本思路[19][20]，也与在线评论场景中使用 RoBERTa 做细粒度情感识别的研究方向一致[12]。",
    "该方法不需要额外训练模型，规则也比较容易检查。它的不足是难以识别隐含属性和跨句表达。":
        "该方法不需要额外训练模型，规则也比较容易检查。它的不足是难以识别隐含属性和跨句表达；从属性值提取和方面级情感研究看，后续若要进一步提高解释粒度，可以引入属性抽取、方面分类或方面-观点联合建模方法[14][16][17]。",
    "这一部分主要用于补充负面评论的解释，不能替代更细的方面级情感分析，但可以帮助用户先看到集中出现的问题。":
        "这一部分主要用于补充负面评论的解释，不能替代更细的方面级情感分析，但可以帮助用户先看到集中出现的问题。也就是说，TextRank 在本文中承担的是“先把高频投诉词找出来”的工程角色，而不是完整的属性级情感模型。",
    "继续补充消融实验，并验证 Focal Loss、Label Smoothing、注意力池化等策略在当前数据集上的效果。":
        "继续补充消融实验，并验证 Focal Loss、Label Smoothing、注意力池化等策略在当前数据集上的效果，同时保留 macro-F1 和类别级 F1 作为主要观察指标[18]。",
    "在功能点分析上尝试方面级联合抽取或属性级情感分析方法，提高对隐式属性和复杂表达的处理能力。":
        "在功能点分析上尝试方面级联合抽取或属性级情感分析方法，提高对隐式属性、方面词和观点词的处理能力[14][15][16][17]。",
    "继续补充合规数据来源，覆盖更多商品品类，提升模型适用范围。":
        "继续补充合规数据来源，覆盖更多商品品类；如果后续接入采集模块，需要结合平台规则和采集框架研究控制频率、去重和异常处理[10]。",
}


EXTRA_AFTER_PARAGRAPH = {
    "近两年的研究进一步说明，这一方向仍在快速演进。胥桂仙等将 BERT 与超图对偶注意力网络结合，尝试增强短文本中的结构信息表达[11]；黄山山基于 RoBERTa 做在线用户评论细粒度情感分析，与本文的模型选型较为接近[12]；金书丞、王嘉梅则从客户评论出发，强调多层语义特征和注意力机制对情感判断的帮助[13]。":
        "从更细的业务维度看，宁秦伟等梳理了面向电子商务的属性值提取研究，说明属性信息抽取是商品信息结构化和评论理解的重要基础[14]；郭之昊等在葡萄电商评论中引入方面词、观点词和情感极性的组合分析，体现了细粒度情感分析的应用价值[15]。国外研究也在从整体情绪转向方面级解释：Davoodi 等用用户评论分析电商平台客户满意度[16]，Dey 和 Jenamani 探索消费者评论中的方面、观点和情感联合建模[17]，Daza 等则从文献计量和系统综述角度总结了电商产品评论情感分析的挑战与趋势[18]。这些工作给本文的启发是，毕业设计不能只做一个分类模型，也不能只做一个后台页面，而要把情感识别、功能点统计、问题挖掘和可视化展示放到同一条流程中。"
}


def backup_docx(path: Path) -> Path:
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    backup_path = path.with_name(f"{path.stem}_before_refs_{timestamp}{path.suffix}")
    shutil.copy2(path, backup_path)
    return backup_path


def insert_paragraph_after(paragraph: Paragraph, text: str, style_name: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style_name:
        new_para.style = style_name
    new_para.add_run(text)
    return new_para


def replace_target_paragraphs(document: Document) -> tuple[int, int]:
    replaced = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        replacement = PARAGRAPH_REPLACEMENTS.get(text)
        if replacement:
            paragraph.text = replacement
            replaced += 1

    inserted = 0
    existing_texts = {p.text.strip() for p in document.paragraphs}
    for trigger, extra_text in EXTRA_AFTER_PARAGRAPH.items():
        if extra_text in existing_texts:
            continue
        for paragraph in document.paragraphs:
            if paragraph.text.strip() == trigger:
                insert_paragraph_after(paragraph, extra_text, paragraph.style.name)
                inserted += 1
                break
    return replaced, inserted


def remove_existing_ref_bookmarks(document: Document) -> None:
    ref_bookmark_ids: set[str] = set()
    for start in list(document.element.xpath(".//w:bookmarkStart")):
        name = start.get(qn("w:name"))
        if name and re.fullmatch(r"ref_\d+", name):
            ref_bookmark_ids.add(start.get(qn("w:id")))
            start.getparent().remove(start)

    for end in list(document.element.xpath(".//w:bookmarkEnd")):
        if end.get(qn("w:id")) in ref_bookmark_ids:
            end.getparent().remove(end)


def next_bookmark_id(document: Document) -> int:
    max_id = 0
    for start in document.element.xpath(".//w:bookmarkStart"):
        value = start.get(qn("w:id"))
        if value and value.isdigit():
            max_id = max(max_id, int(value))
    return max_id + 1


def add_reference_bookmarks(document: Document) -> dict[int, str]:
    remove_existing_ref_bookmarks(document)
    bookmark_id = next_bookmark_id(document)
    found: dict[int, str] = {}

    for paragraph in document.paragraphs:
        match = re.match(r"^\[(\d+)\]", paragraph.text.strip())
        if not match:
            continue
        ref_num = int(match.group(1))
        if not 1 <= ref_num <= REF_COUNT:
            continue

        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(bookmark_id))
        start.set(qn("w:name"), f"ref_{ref_num}")

        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(bookmark_id))

        p = paragraph._p
        insert_at = 1 if len(p) and p[0].tag == qn("w:pPr") else 0
        p.insert(insert_at, start)
        p.insert(insert_at + 1, end)

        found[ref_num] = paragraph.text.strip()
        bookmark_id += 1

    missing = [n for n in range(1, REF_COUNT + 1) if n not in found]
    if missing:
        raise RuntimeError(f"Missing reference paragraphs: {missing}")
    return found


def clear_paragraph_content(paragraph: Paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def make_citation_hyperlink(ref_num: int) -> OxmlElement:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), f"ref_{ref_num}")
    hyperlink.set(qn("w:history"), "1")

    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    run_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    run_pr.append(underline)

    vert_align = OxmlElement("w:vertAlign")
    vert_align.set(qn("w:val"), "superscript")
    run_pr.append(vert_align)

    text = OxmlElement("w:t")
    text.text = f"[{ref_num}]"

    run.append(run_pr)
    run.append(text)
    hyperlink.append(run)
    return hyperlink


def paragraph_is_before_references(index: int, reference_heading_index: int) -> bool:
    return index < reference_heading_index


def rewrite_paragraph_citations(paragraph: Paragraph) -> bool:
    text = paragraph.text
    pattern = re.compile(r"\[(\d{1,2})\]")
    matches = [m for m in pattern.finditer(text) if 1 <= int(m.group(1)) <= REF_COUNT]
    if not matches:
        return False

    clear_paragraph_content(paragraph)
    pos = 0
    for match in matches:
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        paragraph._p.append(make_citation_hyperlink(int(match.group(1))))
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])
    return True


def convert_body_citations(document: Document) -> int:
    reference_heading_index = next(
        (i for i, p in enumerate(document.paragraphs) if p.text.strip() == "参考文献"),
        None,
    )
    if reference_heading_index is None:
        raise RuntimeError("Reference heading not found")

    converted = 0
    for index, paragraph in enumerate(document.paragraphs):
        if not paragraph_is_before_references(index, reference_heading_index):
            continue
        if rewrite_paragraph_citations(paragraph):
            converted += 1
    return converted


def main() -> int:
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)

    backup_path = backup_docx(DOCX_PATH)
    document = Document(str(DOCX_PATH))

    replaced, inserted = replace_target_paragraphs(document)
    references = add_reference_bookmarks(document)
    converted = convert_body_citations(document)

    document.save(str(DOCX_PATH))

    print(f"[OK] backup: {backup_path}")
    print(f"[OK] paragraph replacements: {replaced}, inserted paragraphs: {inserted}")
    print(f"[OK] reference bookmarks rebuilt: {len(references)}")
    print(f"[OK] body paragraphs with converted citations: {converted}")
    print(f"[OK] saved: {DOCX_PATH}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
