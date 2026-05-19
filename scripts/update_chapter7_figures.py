#!/usr/bin/env python3
"""Generate chapter 7 experiment figures and update the thesis DOCX."""

from __future__ import annotations

import csv
import html
import json
import shutil
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "docs/基于NLP技术的电商评论多维可视化分析平台_毕业论文_初稿.docx"
FIGURE_DIR = ROOT / "docs/figures"
SUMMARY_CSV = (
    ROOT
    / "backend/data/experiments/thesis_minimal_20260425_190905/summary/ablation_group_summary.csv"
)
EXPERIMENT_ROOT = ROOT / "backend/data/experiments/thesis_minimal_20260425_190905"
LOG_ROOT = ROOT / "backend/logs/thesis/thesis_minimal_20260425_190905"

GROUPS = [
    ("A0_ce_baseline", "A0", "CE基线"),
    ("A1_ce_fgm", "A1", "CE+FGM"),
    ("A2_ce_fgm_es", "A2", "CE+FGM+早停"),
    ("A3_focal_fgm_es_cw", "A3", "Focal+FGM+早停+权重"),
]
CLASSES = ("negative", "neutral", "positive")
FONT_FAMILY = "PingFang SC, Noto Sans CJK SC, Microsoft YaHei, Arial, sans-serif"


def read_group_summary() -> list[dict[str, float | str]]:
    with SUMMARY_CSV.open(newline="", encoding="utf-8") as f:
        raw_rows = {row["group"]: row for row in csv.DictReader(f)}

    rows: list[dict[str, float | str]] = []
    for group, short_name, description in GROUPS:
        row = raw_rows[group]
        rows.append(
            {
                "group": group,
                "short_name": short_name,
                "description": description,
                "accuracy_mean": float(row["accuracy_mean"]),
                "accuracy_std": float(row["accuracy_std"]),
                "macro_f1_mean": float(row["macro_f1_mean"]),
                "macro_f1_std": float(row["macro_f1_std"]),
            }
        )
    return rows


def parse_classification_reports(log_text: str) -> list[dict[str, object]]:
    reports: list[dict[str, object]] = []
    lines = log_text.splitlines()

    for index, line in enumerate(lines):
        if line.strip() != "分类报告:":
            continue

        class_rows: dict[str, dict[str, float | int]] = {}
        accuracy: float | None = None
        macro_f1: float | None = None

        for report_line in lines[index + 1 : index + 16]:
            parts = report_line.split()
            if not parts:
                continue
            if parts[0] in CLASSES and len(parts) >= 5:
                class_rows[parts[0]] = {
                    "precision": float(parts[1]),
                    "recall": float(parts[2]),
                    "f1": float(parts[3]),
                    "support": int(parts[4]),
                }
            elif parts[0] == "accuracy" and len(parts) >= 2:
                accuracy = float(parts[1])
            elif parts[:2] == ["macro", "avg"] and len(parts) >= 5:
                macro_f1 = float(parts[4])

        if class_rows and macro_f1 is not None:
            reports.append(
                {
                    "class_rows": class_rows,
                    "accuracy": accuracy,
                    "macro_f1": macro_f1,
                }
            )
    return reports


def read_class_f1() -> list[dict[str, float | str]]:
    rows: list[dict[str, float | str]] = []

    for group, short_name, description in GROUPS:
        log_path = LOG_ROOT / f"{group}_seed_42.log"
        summary_path = EXPERIMENT_ROOT / group / "seed_42" / "training_summary.json"
        target_macro_f1 = json.loads(summary_path.read_text(encoding="utf-8"))["best_metrics"][
            "f1_score"
        ]
        reports = parse_classification_reports(log_path.read_text(encoding="utf-8", errors="ignore"))
        if not reports:
            raise RuntimeError(f"No classification report found in {log_path}")

        best_report = min(
            reports,
            key=lambda item: abs(float(item["macro_f1"]) - float(target_macro_f1)),
        )
        class_rows = best_report["class_rows"]
        rows.append(
            {
                "group": group,
                "short_name": short_name,
                "description": description,
                "negative": float(class_rows["negative"]["f1"]),
                "neutral": float(class_rows["neutral"]["f1"]),
                "positive": float(class_rows["positive"]["f1"]),
                "macro_f1": float(best_report["macro_f1"]),
            }
        )
    return rows


def text(
    x: float,
    y: float,
    value: str,
    *,
    size: int = 28,
    fill: str = "#1f2933",
    anchor: str = "middle",
    weight: str = "400",
    rotate: str | None = None,
) -> str:
    transform = f' transform="{rotate}"' if rotate else ""
    return (
        f'<text x="{x:.1f}" y="{y:.1f}" font-family="{FONT_FAMILY}" '
        f'font-size="{size}" fill="{fill}" text-anchor="{anchor}" '
        f'font-weight="{weight}"{transform}>{html.escape(value)}</text>'
    )


def line(x1: float, y1: float, x2: float, y2: float, stroke: str = "#d6dde8", width: int = 2) -> str:
    return (
        f'<line x1="{x1:.1f}" y1="{y1:.1f}" x2="{x2:.1f}" y2="{y2:.1f}" '
        f'stroke="{stroke}" stroke-width="{width}" />'
    )


def rect(
    x: float,
    y: float,
    width: float,
    height: float,
    fill: str,
    *,
    radius: int = 0,
    stroke: str | None = None,
) -> str:
    stroke_attr = f' stroke="{stroke}" stroke-width="2"' if stroke else ""
    return (
        f'<rect x="{x:.1f}" y="{y:.1f}" width="{width:.1f}" height="{height:.1f}" '
        f'rx="{radius}" fill="{fill}"{stroke_attr}/>'
    )


def svg_document(width: int, height: int, body: Iterable[str]) -> str:
    return (
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" '
        f'viewBox="0 0 {width} {height}">\n'
        f'<rect width="{width}" height="{height}" fill="#ffffff"/>\n'
        + "\n".join(body)
        + "\n</svg>\n"
    )


def generate_overall_svg(rows: list[dict[str, float | str]], output_path: Path) -> None:
    width, height = 1800, 1050
    left, right, top, bottom = 130, 90, 140, 220
    plot_width = width - left - right
    plot_height = height - top - bottom
    y_min, y_max = 92.8, 94.0
    ticks = [92.8, 93.0, 93.2, 93.4, 93.6, 93.8, 94.0]
    colors = {"Accuracy": "#2563eb", "Macro-F1": "#f97316"}

    def y(value: float) -> float:
        return top + (y_max - value) / (y_max - y_min) * plot_height

    elements: list[str] = [
        text(width / 2, 62, "总体 Accuracy 与 Macro-F1 对比（A0-A3）", size=40, weight="700"),
        text(width / 2, 105, "三种子均值；误差线表示标准差；纵轴为局部放大（92.8%-94.0%）", size=24, fill="#586475"),
        line(left, top, left, top + plot_height, "#7b8794", 3),
        line(left, top + plot_height, left + plot_width, top + plot_height, "#7b8794", 3),
    ]

    for tick in ticks:
        y_pos = y(tick)
        elements.append(line(left, y_pos, left + plot_width, y_pos, "#e6ebf2", 2))
        elements.append(text(left - 18, y_pos + 8, f"{tick:.1f}%", size=22, fill="#52606d", anchor="end"))

    group_width = plot_width / len(rows)
    bar_width = 78
    bar_gap = 22
    baseline_y = y(y_min)

    for index, row in enumerate(rows):
        group_center = left + group_width * index + group_width / 2
        label = f"{row['short_name']}（{row['description']}）"
        elements.append(text(group_center, baseline_y + 50, str(label), size=22, fill="#323f4b"))

        specs = [
            ("Accuracy", float(row["accuracy_mean"]) * 100, float(row["accuracy_std"]) * 100),
            ("Macro-F1", float(row["macro_f1_mean"]) * 100, float(row["macro_f1_std"]) * 100),
        ]
        start_x = group_center - bar_width - bar_gap / 2
        for metric_index, (metric, mean_value, std_value) in enumerate(specs):
            x_pos = start_x + metric_index * (bar_width + bar_gap)
            y_pos = y(mean_value)
            bar_height = baseline_y - y_pos
            color = colors[metric]
            elements.append(rect(x_pos, y_pos, bar_width, bar_height, color, radius=5))
            elements.append(text(x_pos + bar_width / 2, y_pos - 14, f"{mean_value:.2f}%", size=22, fill="#1f2933", weight="600"))

            error_top = y(mean_value + std_value)
            error_bottom = y(mean_value - std_value)
            error_x = x_pos + bar_width / 2
            elements.append(line(error_x, error_top, error_x, error_bottom, "#111827", 2))
            elements.append(line(error_x - 16, error_top, error_x + 16, error_top, "#111827", 2))
            elements.append(line(error_x - 16, error_bottom, error_x + 16, error_bottom, "#111827", 2))

    legend_x = width - 430
    legend_y = 145
    for offset, metric in enumerate(("Accuracy", "Macro-F1")):
        y_pos = legend_y + offset * 42
        elements.append(rect(legend_x, y_pos - 23, 34, 22, colors[metric], radius=3))
        elements.append(text(legend_x + 48, y_pos - 4, metric, size=24, anchor="start", fill="#323f4b"))

    elements.append(text(left + plot_width / 2, height - 32, "实验组", size=24, fill="#52606d"))
    elements.append(
        text(
            34,
            top + plot_height / 2,
            "指标值",
            size=24,
            fill="#52606d",
            rotate=f"rotate(-90 34 {top + plot_height / 2:.1f})",
        )
    )

    output_path.write_text(svg_document(width, height, elements), encoding="utf-8")


def generate_class_f1_svg(rows: list[dict[str, float | str]], output_path: Path) -> None:
    width, height = 1800, 1050
    left, right, top, bottom = 130, 90, 150, 230
    plot_width = width - left - right
    plot_height = height - top - bottom
    y_min, y_max = 0.910, 0.945
    ticks = [0.910, 0.915, 0.920, 0.925, 0.930, 0.935, 0.940, 0.945]
    colors = {
        "A0": "#6b7280",
        "A1": "#2563eb",
        "A2": "#16a34a",
        "A3": "#dc2626",
    }

    def y(value: float) -> float:
        return top + (y_max - value) / (y_max - y_min) * plot_height

    elements: list[str] = [
        text(width / 2, 64, "类别级 F1 对比（seed=42 最佳轮次）", size=40, weight="700"),
        text(width / 2, 108, "纵轴为局部放大（0.910-0.945），用于观察接近数值之间的差异", size=24, fill="#586475"),
        line(left, top, left, top + plot_height, "#7b8794", 3),
        line(left, top + plot_height, left + plot_width, top + plot_height, "#7b8794", 3),
    ]

    for tick in ticks:
        y_pos = y(tick)
        elements.append(line(left, y_pos, left + plot_width, y_pos, "#e6ebf2", 2))
        elements.append(text(left - 18, y_pos + 8, f"{tick:.3f}", size=22, fill="#52606d", anchor="end"))

    category_width = plot_width / len(CLASSES)
    bar_width = 62
    bar_gap = 18
    baseline_y = y(y_min)
    class_names = {"negative": "negative", "neutral": "neutral", "positive": "positive"}

    for class_index, class_name in enumerate(CLASSES):
        category_center = left + category_width * class_index + category_width / 2
        total_bar_width = len(rows) * bar_width + (len(rows) - 1) * bar_gap
        start_x = category_center - total_bar_width / 2
        elements.append(text(category_center, baseline_y + 56, class_names[class_name], size=26, fill="#323f4b", weight="600"))

        for row_index, row in enumerate(rows):
            value = float(row[class_name])
            x_pos = start_x + row_index * (bar_width + bar_gap)
            y_pos = y(value)
            bar_height = baseline_y - y_pos
            color = colors[str(row["short_name"])]
            elements.append(rect(x_pos, y_pos, bar_width, bar_height, color, radius=5))
            elements.append(text(x_pos + bar_width / 2, y_pos - 12, f"{value:.4f}", size=20, fill="#1f2933", weight="600"))

    legend_x = width - 610
    legend_y = 150
    for index, row in enumerate(rows):
        x_pos = legend_x + index * 150
        short_name = str(row["short_name"])
        elements.append(rect(x_pos, legend_y - 24, 34, 22, colors[short_name], radius=3))
        elements.append(text(x_pos + 45, legend_y - 5, f"{short_name}", size=24, anchor="start", fill="#323f4b"))

    elements.append(text(left + plot_width / 2, height - 32, "类别", size=24, fill="#52606d"))
    elements.append(
        text(
            34,
            top + plot_height / 2,
            "F1 值",
            size=24,
            fill="#52606d",
            rotate=f"rotate(-90 34 {top + plot_height / 2:.1f})",
        )
    )

    output_path.write_text(svg_document(width, height, elements), encoding="utf-8")


def render_png(svg_path: Path, png_path: Path) -> None:
    subprocess.run(
        ["rsvg-convert", "-w", "1800", "-h", "1050", "-o", str(png_path), str(svg_path)],
        check=True,
    )


def set_paragraph_spacing(paragraph, *, before: int = 0, after: int = 6, line: float = 1.2) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(before)
    paragraph_format.space_after = Pt(after)
    paragraph_format.line_spacing = line


def block_element(block):
    return block._p if hasattr(block, "_p") else block._tbl


def insert_paragraph_after(document: Document, cursor, text_value: str = "", style: str | None = None):
    paragraph = document.add_paragraph()
    if style:
        paragraph.style = style
    if text_value:
        paragraph.add_run(text_value)
    block_element(cursor).addnext(paragraph._p)
    if style in {"Normal", None}:
        set_paragraph_spacing(paragraph)
    return paragraph


def insert_picture_after(document: Document, cursor, image_path: Path):
    paragraph = insert_paragraph_after(document, cursor)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(6.35))
    return paragraph


def insert_table_after(document: Document, cursor, rows: list[list[str]]):
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for row_index, row_values in enumerate(rows):
        for col_index, value in enumerate(row_values):
            cell = table.cell(row_index, col_index)
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_spacing(paragraph, after=0, line=1.0)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    if row_index == 0:
                        run.bold = True
    block_element(cursor).addnext(table._tbl)
    return table


def find_paragraph(document: Document, text_value: str):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text_value:
            return paragraph
    raise ValueError(f"Paragraph not found: {text_value}")


def remove_blocks_between(document: Document, start_paragraph, end_paragraph) -> None:
    body = document.element.body
    children = list(body.iterchildren())
    start_index = children.index(start_paragraph._p)
    end_index = children.index(end_paragraph._p)
    for child in children[start_index + 1 : end_index]:
        body.remove(child)


def enable_field_update_on_open(document: Document) -> None:
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def update_docx(group_rows, class_rows, figure_1_png: Path, figure_2_png: Path) -> Path:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_path = DOCX_PATH.with_name(f"{DOCX_PATH.stem}_backup_{timestamp}.docx")
    shutil.copy2(DOCX_PATH, backup_path)

    document = Document(str(DOCX_PATH))
    section_heading = find_paragraph(document, "7.4 模型训练结果分析")
    next_heading = find_paragraph(document, "7.5 本章小结")
    chapter8_heading = find_paragraph(document, "第八章 总结与展望")
    remove_blocks_between(document, section_heading, next_heading)
    remove_blocks_between(document, next_heading, chapter8_heading)

    cursor = section_heading
    cursor = insert_paragraph_after(document, cursor, "7.4.1 总体结果分析", "Heading 3")
    cursor = insert_paragraph_after(
        document,
        cursor,
        "本节使用同一训练集 train_balanced_full.csv 上的四组消融实验结果进行比较。A0 表示 RoBERTa 加交叉熵损失的基础模型；A1 在 A0 基础上加入 FGM 对抗训练；A2 在 A1 基础上加入 Early Stopping；A3 使用 Focal Loss、FGM、Early Stopping 和类别权重。这样设置的目的，是在其他主要训练参数保持一致的情况下，观察不同训练策略对三分类结果的影响。",
        "Normal",
    )
    cursor = insert_paragraph_after(
        document,
        cursor,
        "图 7-1 采用三个随机种子的均值，并把标准差作为误差线。由于四组结果本身比较接近，图中对纵轴做了局部放大，读图时应结合柱顶标注的具体数值判断差异。",
        "Normal",
    )
    cursor = insert_picture_after(document, cursor, figure_1_png)
    cursor = insert_paragraph_after(document, cursor, "图7-1总体 Accuracy 与 Macro-F1 对比柱状图", "Normal")
    cursor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cursor = insert_paragraph_after(
        document,
        cursor,
        "从总体指标看，A1 和 A2 相比 A0 都有提升，说明 FGM 对抗训练对当前评论分类任务有帮助。A2 的 Accuracy 均值为 93.77%，Macro-F1 均值为 93.57%，是四组中最高的一组；同时 A2 的 Macro-F1 标准差为 0.0864%，也低于其他组。A3 加入 Focal Loss 和类别权重后没有继续超过 A2，说明在当前较平衡的数据集上，更复杂的损失设置不一定带来更好的平均表现。",
        "Normal",
    )

    cursor = insert_paragraph_after(document, cursor, "7.4.2 类别级结果分析", "Heading 3")
    cursor = insert_paragraph_after(
        document,
        cursor,
        "类别级结果取各实验组 seed=42 运行中最佳 Macro-F1 轮次的分类报告。这里不只看总体准确率，是因为电商评论三分类任务中，中性评论更容易受到正负表达混杂、语气较弱等因素影响，单独列出类别 F1 更容易看出模型的短板和变化。",
        "Normal",
    )
    cursor = insert_paragraph_after(document, cursor, "表7-3 A0-A3 类别级 F1 对比", "Normal")

    table_rows = [["实验组", "Negative F1", "Neutral F1", "Positive F1"]]
    for row in class_rows:
        table_rows.append(
            [
                f"{row['short_name']}（{row['description']}）",
                f"{float(row['negative']):.4f}",
                f"{float(row['neutral']):.4f}",
                f"{float(row['positive']):.4f}",
            ]
        )
    cursor = insert_table_after(document, cursor, table_rows)
    cursor = insert_picture_after(document, cursor, figure_2_png)
    cursor = insert_paragraph_after(document, cursor, "图7-2类别级 F1 对比图", "Normal")
    cursor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cursor = insert_paragraph_after(
        document,
        cursor,
        "从表 7-3 和图 7-2 可以看到，A1、A2 相比 A0 的提升主要体现在 neutral 类和 positive 类上。A2 的 neutral F1 为 0.9262，略高于 A1 的 0.9258；negative 类上 A1 与 A2 的差距很小。A3 的三个类别 F1 都保持在较高水平，但没有成为最优结果，因此本文最终更倾向于采用 A2 作为后续模型配置。",
        "Normal",
    )

    cursor = insert_paragraph_after(document, cursor, "7.4.3 当前实验边界", "Heading 3")
    cursor = insert_paragraph_after(
        document,
        cursor,
        "本节引用的总体指标来自 ablation_group_summary.csv，类别级指标来自对应训练日志中的分类报告。由于类别级 F1 没有单独汇总三种子的均值文件，图 7-2 采用 seed=42 的代表性运行结果；因此它主要用于观察不同类别的相对变化，不把它解释为所有随机种子下的类别均值。后续如果继续完善实验，可以把每个种子的类别级分类报告统一导出，再补充类别级均值和标准差。",
        "Normal",
    )

    insert_paragraph_after(
        document,
        next_heading,
        "本章整理了测试代码、平台功能证据和模型实验结果。结合现有代码、PDF 产物、消融实验汇总和训练日志，平台主要功能已经能够连贯运行；在模型部分，A2（CE + FGM + Early Stopping）在总体指标和稳定性上表现更合适，因此可作为当前论文实验中的主要配置。",
        "Normal",
    )

    enable_field_update_on_open(document)

    try:
        document.save(str(DOCX_PATH))
        output_path = DOCX_PATH
    except PermissionError:
        output_path = DOCX_PATH.with_name(f"{DOCX_PATH.stem}_第七章图表更新版_{timestamp}.docx")
        document.save(str(output_path))

    print(f"[OK] backup: {backup_path}")
    print(f"[OK] docx: {output_path}")
    return output_path


def main() -> int:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)

    group_rows = read_group_summary()
    class_rows = read_class_f1()

    figure_1_svg = FIGURE_DIR / "figure_7_1_overall_accuracy_macro_f1.svg"
    figure_1_png = FIGURE_DIR / "figure_7_1_overall_accuracy_macro_f1.png"
    figure_2_svg = FIGURE_DIR / "figure_7_2_class_f1.svg"
    figure_2_png = FIGURE_DIR / "figure_7_2_class_f1.png"

    generate_overall_svg(group_rows, figure_1_svg)
    generate_class_f1_svg(class_rows, figure_2_svg)
    render_png(figure_1_svg, figure_1_png)
    render_png(figure_2_svg, figure_2_png)
    update_docx(group_rows, class_rows, figure_1_png, figure_2_png)

    print(f"[OK] figure: {figure_1_png}")
    print(f"[OK] figure: {figure_2_png}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
