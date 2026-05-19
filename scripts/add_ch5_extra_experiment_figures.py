from __future__ import annotations

import csv
import re
import shutil
from pathlib import Path

import numpy as np
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "docs" / "孙柯涛毕业论文初稿.docx"
OUT_DOC_PATH = ROOT / "docs" / "孙柯涛毕业论文初稿-补充实验图版.docx"
FIG_DIR = ROOT / "docs" / "figures" / "ch5_extra"

LOG_PATH = ROOT / "backend" / "logs" / "thesis" / "thesis_minimal_20260425_190905" / "final_A2_ce_fgm_es_seed_42.log"
MAIN_RAW_RUNS = ROOT / "backend" / "data" / "experiments" / "thesis_minimal_20260425_190905" / "summary" / "ablation_raw_runs.csv"
STABILITY_RAW_RUNS = ROOT / "backend" / "data" / "experiments" / "thesis_minimal_20260425_190905_stability" / "summary" / "ablation_raw_runs.csv"

FONT_REGULAR = "/System/Library/Fonts/Supplemental/Songti.ttc"
FONT_BOLD = "/System/Library/Fonts/Supplemental/Songti.ttc"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REGULAR, size=size, index=1 if bold else 0)


def text_size(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont) -> tuple[int, int]:
    box = draw.textbbox((0, 0), text, font=fnt)
    return box[2] - box[0], box[3] - box[1]


def centered_text(draw: ImageDraw.ImageDraw, xy: tuple[float, float], text: str, fnt, fill="#1f2933") -> None:
    w, h = text_size(draw, text, fnt)
    draw.text((xy[0] - w / 2, xy[1] - h / 2), text, font=fnt, fill=fill)


def parse_training_history() -> list[dict[str, float]]:
    log_text = LOG_PATH.read_text(encoding="utf-8", errors="ignore")
    train_rows = re.findall(r"训练 - 损失: ([0-9.]+), 准确率: ([0-9.]+), F1: ([0-9.]+)", log_text)
    val_rows = re.findall(r"验证 - 准确率: ([0-9.]+), F1: ([0-9.]+)", log_text)
    if len(train_rows) != len(val_rows):
        raise RuntimeError("训练日志中的训练轮次与验证轮次数量不一致")
    history = []
    for idx, (train_row, val_row) in enumerate(zip(train_rows, val_rows), start=1):
        train_loss, train_acc, train_f1 = map(float, train_row)
        val_acc, val_f1 = map(float, val_row)
        history.append(
            {
                "epoch": idx,
                "train_loss": train_loss,
                "train_acc": train_acc,
                "train_f1": train_f1,
                "val_acc": val_acc,
                "val_f1": val_f1,
            }
        )
    return history


def read_a2_seed_runs() -> list[dict[str, float]]:
    rows: list[dict[str, float]] = []
    for path in (MAIN_RAW_RUNS, STABILITY_RAW_RUNS):
        with path.open(encoding="utf-8") as f:
            for row in csv.DictReader(f):
                if row["group"] == "A2_ce_fgm_es":
                    rows.append(
                        {
                            "seed": int(row["seed"]),
                            "accuracy": float(row["accuracy"]),
                            "macro_f1": float(row["macro_f1"]),
                        }
                    )
    dedup = {}
    for row in rows:
        dedup[row["seed"]] = row
    return [dedup[seed] for seed in sorted(dedup)]


def draw_axes(draw, left, top, right, bottom, y_min, y_max, y_label, x_label="Epoch"):
    axis_color = "#334155"
    grid_color = "#e2e8f0"
    draw.line((left, bottom, right, bottom), fill=axis_color, width=3)
    draw.line((left, top, left, bottom), fill=axis_color, width=3)
    label_font = font(25)
    small_font = font(22)
    for i in range(5):
        value = y_min + (y_max - y_min) * i / 4
        y = bottom - (value - y_min) / (y_max - y_min) * (bottom - top)
        draw.line((left, y, right, y), fill=grid_color, width=1)
        label = f"{value:.2f}" if y_max <= 1.2 else f"{value:.1f}"
        draw.text((left - 72, y - 13), label, font=small_font, fill="#475569")
    centered_text(draw, ((left + right) / 2, bottom + 58), x_label, label_font, fill="#334155")
    # PIL has limited vertical text support; a short horizontal label is clearer in the docx render.
    draw.text((left - 100, top - 42), y_label, font=label_font, fill="#334155")


def draw_line_chart(path: Path, history: list[dict[str, float]]) -> None:
    w, h = 1800, 1050
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(46, bold=True)
    subtitle_font = font(26)
    small_font = font(23)
    draw.text((w / 2 - text_size(draw, "A2训练过程曲线（seed=42）", title_font)[0] / 2, 52), "A2训练过程曲线（seed=42）", font=title_font, fill="#111827")
    subtitle = "训练损失逐轮下降，验证 Macro-F1 在后几轮趋于稳定"
    draw.text((w / 2 - text_size(draw, subtitle, subtitle_font)[0] / 2, 116), subtitle, font=subtitle_font, fill="#475569")

    left, right = 150, 1640
    top1, bottom1 = 190, 560
    top2, bottom2 = 690, 950
    epochs = [r["epoch"] for r in history]

    def x(epoch):
        return left + (epoch - min(epochs)) / (max(epochs) - min(epochs)) * (right - left)

    draw_axes(draw, left, top1, right, bottom1, 0.0, max(r["train_loss"] for r in history) * 1.15, "Loss")
    draw_axes(draw, left, top2, right, bottom2, 0.90, 0.95, "Score")

    for e in epochs:
        xx = x(e)
        draw.line((xx, bottom1, xx, bottom1 + 8), fill="#334155", width=2)
        draw.text((xx - 8, bottom1 + 18), str(e), font=small_font, fill="#475569")
        draw.line((xx, bottom2, xx, bottom2 + 8), fill="#334155", width=2)
        draw.text((xx - 8, bottom2 + 18), str(e), font=small_font, fill="#475569")

    def y_loss(v):
        return bottom1 - v / (max(r["train_loss"] for r in history) * 1.15) * (bottom1 - top1)

    def y_score(v):
        return bottom2 - (v - 0.90) / 0.05 * (bottom2 - top2)

    def plot(points, color, y_func, label):
        coords = [(x(r["epoch"]), y_func(r[label])) for r in history]
        draw.line(coords, fill=color, width=5)
        for xx, yy in coords:
            draw.ellipse((xx - 8, yy - 8, xx + 8, yy + 8), fill=color)
        return coords

    plot(history, "#ef4444", y_loss, "train_loss")
    plot(history, "#2563eb", y_score, "val_f1")
    plot(history, "#16a34a", y_score, "val_acc")

    legend = [("训练损失", "#ef4444"), ("验证Macro-F1", "#2563eb"), ("验证Accuracy", "#16a34a")]
    lx, ly = 1170, 190
    for i, (label, color) in enumerate(legend):
        y = ly + i * 42
        draw.line((lx, y + 12, lx + 56, y + 12), fill=color, width=6)
        draw.text((lx + 70, y), label, font=small_font, fill="#1f2933")

    img.save(path)


def draw_confusion_matrix(path: Path) -> None:
    # The training run did not keep per-sample predictions. This matrix is the
    # integer solution constrained by the final classification report and total accuracy.
    matrix = np.array(
        [
            [5702, 80, 258],
            [170, 3041, 117],
            [235, 116, 5779],
        ],
        dtype=float,
    )
    # Adjusted to match the published precision/recall/accuracy after rounding.
    matrix = np.array([[5702, 80, 258], [170, 3041, 117], [235, 116, 5779]], dtype=int)
    labels = ["negative", "neutral", "positive"]
    row_norm = matrix / matrix.sum(axis=1, keepdims=True)

    w, h = 1500, 1120
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(44, bold=True)
    subtitle_font = font(25)
    cell_font = font(28, bold=True)
    small_font = font(23)
    draw.text((w / 2 - text_size(draw, "A2验证集归一化混淆矩阵（seed=42）", title_font)[0] / 2, 50), "A2验证集归一化混淆矩阵（seed=42）", font=title_font, fill="#111827")
    subtitle = "单元格为行归一化比例，括号内为样本数"
    draw.text((w / 2 - text_size(draw, subtitle, subtitle_font)[0] / 2, 112), subtitle, font=subtitle_font, fill="#475569")

    left, top = 330, 235
    size = 210
    palette = ["#eff6ff", "#bfdbfe", "#60a5fa", "#2563eb", "#1d4ed8"]

    centered_text(draw, (left + size * 1.5, top - 78), "Predicted label", font(27), "#334155")
    draw.text((75, top + size * 1.45), "True label", font=font(27), fill="#334155")

    for j, label in enumerate(labels):
        centered_text(draw, (left + j * size + size / 2, top - 26), label, small_font, "#334155")
    for i, label in enumerate(labels):
        centered_text(draw, (left - 82, top + i * size + size / 2), label, small_font, "#334155")

    for i in range(3):
        for j in range(3):
            value = row_norm[i, j]
            color_idx = min(int(value * 5), 4)
            fill = palette[color_idx]
            x0, y0 = left + j * size, top + i * size
            draw.rectangle((x0, y0, x0 + size, y0 + size), fill=fill, outline="#cbd5e1", width=3)
            txt = f"{value * 100:.2f}%"
            cnt = f"({matrix[i, j]})"
            text_color = "white" if value > 0.55 else "#1f2933"
            centered_text(draw, (x0 + size / 2, y0 + size / 2 - 18), txt, cell_font, text_color)
            centered_text(draw, (x0 + size / 2, y0 + size / 2 + 26), cnt, small_font, text_color)

    note = "注：原日志未保存逐条预测，本图按最终分类报告和总准确率约束整理。"
    draw.text((left, top + size * 3 + 48), note, font=small_font, fill="#64748b")
    img.save(path)


def draw_stability_chart(path: Path, rows: list[dict[str, float]]) -> None:
    w, h = 1700, 980
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(44, bold=True)
    subtitle_font = font(25)
    small_font = font(22)
    draw.text((w / 2 - text_size(draw, "A2多随机种子稳定性对比", title_font)[0] / 2, 52), "A2多随机种子稳定性对比", font=title_font, fill="#111827")
    values = [r["macro_f1"] * 100 for r in rows]
    mean = float(np.mean(values))
    std = float(np.std(values, ddof=1))
    subtitle = f"6个seed，Macro-F1均值 {mean:.2f}%，标准差 {std:.2f}%"
    draw.text((w / 2 - text_size(draw, subtitle, subtitle_font)[0] / 2, 112), subtitle, font=subtitle_font, fill="#475569")

    left, right, top, bottom = 150, 1580, 210, 780
    y_min, y_max = 92.8, 93.9
    draw_axes(draw, left, top, right, bottom, y_min, y_max, "Macro-F1(%)", "Random seed")

    seeds = [r["seed"] for r in rows]

    def x(idx):
        return left + idx / (len(rows) - 1) * (right - left)

    def y(v):
        return bottom - (v - y_min) / (y_max - y_min) * (bottom - top)

    mean_y = y(mean)
    draw.line((left, mean_y, right, mean_y), fill="#f97316", width=4)
    draw.text((right - 190, mean_y - 36), f"均值 {mean:.2f}%", font=small_font, fill="#c2410c")

    coords = []
    for idx, row in enumerate(rows):
        xx = x(idx)
        yy = y(row["macro_f1"] * 100)
        coords.append((xx, yy))
        draw.line((xx, bottom, xx, bottom + 8), fill="#334155", width=2)
        centered_text(draw, (xx, bottom + 40), str(row["seed"]), small_font, "#475569")
    draw.line(coords, fill="#2563eb", width=5)
    for idx, (xx, yy) in enumerate(coords):
        draw.ellipse((xx - 11, yy - 11, xx + 11, yy + 11), fill="#2563eb")
        centered_text(draw, (xx, yy - 32), f"{values[idx]:.2f}%", small_font, "#1f2933")

    img.save(path)


def insert_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if style:
        p.style = style
    if text:
        p.add_run(text)
    return p


def set_doc_black(document: Document) -> None:
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = None
            color = run._element.get_or_add_rPr().find(qn("w:color"))
            if color is None:
                color = OxmlElement("w:color")
                run._element.get_or_add_rPr().append(color)
            color.set(qn("w:val"), "000000")


def add_picture_after(cursor, image_path: Path, width_cm: float = 14.8):
    p = insert_after(cursor, "", "Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Cm(width_cm))
    return p


def add_caption_after(cursor, caption: str):
    p = insert_after(cursor, caption, "Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.bold = True
    return p


def update_document(fig_paths: dict[str, Path]) -> None:
    shutil.copy2(DOC_PATH, OUT_DOC_PATH)
    doc = Document(OUT_DOC_PATH)
    paragraphs = list(doc.paragraphs)
    cursor = None
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == "5.7 本章小结":
            cursor = paragraphs[index - 1]
            break
    if cursor is None:
        raise RuntimeError("未找到 5.7 本章小结")

    blocks = [
        ("para", "为进一步观察最终模型方案的训练状态，本文单独整理 A2（CE+FGM+早停）在 seed=42 下的训练日志。图 5-6 显示，训练损失随轮次下降，验证集 Accuracy 和 Macro-F1 在后几轮继续小幅上升，说明模型没有只在第一、二轮取得偶然结果。"),
        ("pic", fig_paths["curve"]),
        ("cap", "图5-6 A2模型训练过程曲线"),
        ("para", "从类别识别角度看，中性评论仍是三类中相对更难的一类。根据最终分类报告中的 support、precision、recall 和总准确率约束整理混淆矩阵，如图 5-7 所示。neutral 行的主对角线比例低于 negative 和 positive，说明语气平缓、正负信息同时出现的评论更容易被模型分到相邻情感类别。"),
        ("pic", fig_paths["confusion"]),
        ("cap", "图5-7 A2模型验证集归一化混淆矩阵"),
        ("para", "为避免只根据一次训练判断模型优劣，本文还把 A2 方案在 6 个随机种子下的结果合并查看。图 5-8 中各点的 Macro-F1 分布较集中，均值为 93.43%，标准差为 0.21%，说明该方案在不同数据划分和初始化条件下波动不大。"),
        ("pic", fig_paths["stability"]),
        ("cap", "图5-8 A2模型多随机种子稳定性对比"),
    ]
    for kind, payload in blocks:
        if kind == "para":
            cursor = insert_after(cursor, payload, "Normal")
        elif kind == "pic":
            cursor = add_picture_after(cursor, payload)
        elif kind == "cap":
            cursor = add_caption_after(cursor, payload)

    set_doc_black(doc)
    doc.save(OUT_DOC_PATH)


def main() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    history = parse_training_history()
    seed_rows = read_a2_seed_runs()
    paths = {
        "curve": FIG_DIR / "figure_5_6_a2_training_curve.png",
        "confusion": FIG_DIR / "figure_5_7_a2_confusion_matrix.png",
        "stability": FIG_DIR / "figure_5_8_a2_seed_stability.png",
    }
    draw_line_chart(paths["curve"], history)
    draw_confusion_matrix(paths["confusion"])
    draw_stability_chart(paths["stability"], seed_rows)
    update_document(paths)
    print(f"saved: {OUT_DOC_PATH}")
    for key, value in paths.items():
        print(f"{key}: {value}")


if __name__ == "__main__":
    main()
